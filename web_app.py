import streamlit as st
import pandas as pd
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# 页面基础配置
st.set_page_config(page_title="开业花篮条幅自动生成器", page_icon="💐", layout="centered")

st.title("💐 开业花篮条幅生成器")
st.markdown("上传 Excel 名单，填入公司名称与祝福语，一键生成符合格式要求的打印版 Word 文档！")
st.divider()

# --- 1. 输入区设置 ---
target_company = st.text_input(
    "1. 被祝福公司完整全称（抬头上联）：", 
    value="深圳市瑞康光联科技有限公司",
    placeholder="请输入被祝福公司的完整正式名称"
)

uploaded_file = st.file_uploader(
    "2. 上传祝福公司名单 Excel 文件（.xlsx / .xls）：", 
    type=["xlsx", "xls"]
)

blessings_text = st.text_area(
    "3. 轮换祝福语列表（每行输入一个，自动交替轮换）：",
    value="生意兴隆\n财源广进\n大展宏图\n前程似锦",
    height=120
)

# --- 字体设置函数 ---
def set_run_font(run, font_name, size_pt, bold=True):
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

# --- 生成 Word 文档核心函数 ---
def generate_docx(df, target_company, blessings):
    # 1. 查找名称列
    name_col = None
    for col in df.columns:
        col_str = str(col).strip()
        if any(k in col_str for k in ["公司", "名称", "单位", "个人", "姓名", "客户"]):
            name_col = col
            break

    if not name_col:
        # 若未按关键字匹配到，根据常见表格结构推断
        if len(df.columns) >= 2 and any(k in str(df.columns[0]).lower() for k in ["id", "no", "序"]):
            name_col = df.columns[1]
        else:
            name_col = df.columns[0]

    # 2. 数量列识别：优先取第 3 列（索引 2），不足 3 列时通过表头关键词识别
    qty_col = None
    if len(df.columns) >= 3:
        qty_col = df.columns[2]
    else:
        for col in df.columns:
            col_str = str(col).strip()
            if any(k in col_str for k in ["数", "份", "对", "个", "数量"]):
                qty_col = col
                break

    # 清理数据
    df_clean = df.dropna(subset=[name_col]).copy()
    df_clean = df_clean[~df_clean[name_col].astype(str).str.contains("合计|汇总|总计", na=False)]

    doc = docx.Document()
    
    # 页面属性设置（A4 横向）
    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = Pt(841.9)
    section.page_height = Pt(595.3)
    section.top_margin = Pt(41.95)
    section.bottom_margin = Pt(33.45)
    section.left_margin = Pt(29.5)
    section.right_margin = Pt(40.8)

    card_count = 0
    blessing_idx = 0

    # 抬头公司名称字号动态自适应计算（避免过长折行挤压下文）
    t_len = len(target_company)
    if t_len <= 12:
        target_font_size = 43
    elif t_len <= 16:
        target_font_size = 36
    elif t_len <= 20:
        target_font_size = 30
    else:
        target_font_size = max(18, int(600 / t_len))

    for _, row in df_clean.iterrows():
        sender_name = str(row[name_col]).strip()
        
        # 解析数量（支持浮点型数字转换及默认兜底）
        qty = 1
        if qty_col and pd.notna(row[qty_col]):
            try:
                qty = int(float(row[qty_col]))
                if qty <= 0:
                    qty = 1
            except (ValueError, TypeError):
                qty = 1

        # 按照数量循环生成条幅
        for _ in range(qty):
            card_count += 1

            # 轮换选取祝福语（同公司多张条幅自动轮换下一条不同祝福语）
            blessing = blessings[blessing_idx % len(blessings)]
            blessing_idx += 1

            # 1. 抬头（左对齐，单张条幅起始页分页控制）
            p0 = doc.add_paragraph()
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            p0.paragraph_format.line_spacing = 1.0
            
            # 从第 2 张条幅开始强制前置分页，确保每张条幅严格独占 1 页
            if card_count > 1:
                p0.paragraph_format.page_break_before = True

            r0_1 = p0.add_run("祝")
            set_run_font(r0_1, "宋体", 65, True)
            r0_2 = p0.add_run("：")
            set_run_font(r0_2, "宋体", 56, True)
            r0_3 = p0.add_run(target_company)
            set_run_font(r0_3, "宋体", target_font_size, True)

            # 2. 祝福语（居中对齐，根据文字长度自适应字号）
            p1 = doc.add_paragraph()
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            p1.paragraph_format.line_spacing = 1.0

            b_len = len(blessing)
            if b_len <= 4:
                blessing_font_size = 192
            elif b_len <= 6:
                blessing_font_size = 130
            elif b_len <= 8:
                blessing_font_size = 96
            else:
                blessing_font_size = max(60, int(750 / b_len))

            r1 = p1.add_run(blessing)
            set_run_font(r1, "宋体", blessing_font_size, True)

            # 3. 落款（右对齐，根据落款公司字数自适应字号，确保不折行溢出）
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.line_spacing = 1.0

            s_len = len(sender_name)
            if s_len <= 8:
                sender_font_size = 48
            elif s_len <= 12:
                sender_font_size = 42
            elif s_len <= 16:
                sender_font_size = 34
            elif s_len <= 20:
                sender_font_size = 28
            elif s_len <= 25:
                sender_font_size = 22
            else:
                sender_font_size = max(16, int(560 / s_len))

            r2_1 = p2.add_run(sender_name)
            set_run_font(r2_1, "宋体", sender_font_size, True)

            r2_2 = p2.add_run(" ")
            set_run_font(r2_2, "宋体", 56, True)

            r2_3 = p2.add_run("贺")
            set_run_font(r2_3, "方正粗黑宋简体", 96, True)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer, card_count

# --- 2. 生成与下载按钮 ---
st.divider()
if st.button("🚀 生成 Word 条幅", type="primary", use_container_width=True):
    if not target_company.strip():
        st.warning("⚠️ 请输入【被祝福公司】的名称！")
    elif not uploaded_file:
        st.warning("⚠️ 请先上传 Excel 文件！")
    else:
        blessings_list = [b.strip() for b in blessings_text.split("\n") if b.strip()]
        if not blessings_list:
            st.warning("⚠️ 请至少输入一条祝福语！")
        else:
            try:
                df = pd.read_excel(uploaded_file)
                result = generate_docx(df, target_company.strip(), blessings_list)
                if result:
                    doc_buffer, count = result
                    st.success(f"🎉 成功生成！共包含 {count} 张条幅。")
                    st.download_button(
                        label="⬇️ 点击下载 Word 文件 (.docx)",
                        data=doc_buffer,
                        file_name=f"{target_company.strip()}_开业花篮条幅.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            except Exception as e:
                st.error(f"❌ 处理发生错误: {e}")
